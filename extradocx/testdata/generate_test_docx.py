"""
Generate a rich test.docx with:
  - Cover page (title, subtitle)
  - Table of Contents section
  - 6 chapters with h1/h2/h3 headings
  - Body paragraphs, bold/italic text
  - Bullet and numbered lists
  - A table per chapter
  - Code-style paragraph
  - Links
  - 20+ pages total
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---------------------------------------------------------------------------
# Cover Page
# ---------------------------------------------------------------------------
title = doc.add_heading("Comprehensive Software Engineering Report", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph("A Practical Guide to Modern Software Development Practices")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.runs[0]
run.italic = True
run.font.size = Pt(14)

doc.add_paragraph(
    "Author: Jane Smith\nDate: 2025-04-08\nVersion: 3.1"
).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# ---------------------------------------------------------------------------
# Helper: add a lorem ipsum paragraph  
# ---------------------------------------------------------------------------
LOREM = (
    "Lorem ipsum dolor sit amet, consectetur adipiscing elit. "
    "Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. "
    "Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris "
    "nisi ut aliquip ex ea commodo consequat. Duis aute irure dolor in "
    "reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla "
    "pariatur. Excepteur sint occaecat cupidatat non proident, sunt in "
    "culpa qui officia deserunt mollit anim id est laborum. "
)

def lorem(doc, n=1):
    for _ in range(n):
        doc.add_paragraph(LOREM)

def add_bullet_list(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')

def add_numbered_list(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style='List Number')

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph("")

def add_mixed_paragraph(doc, text):
    """Paragraph with bold and italic runs."""
    p = doc.add_paragraph()
    p.add_run("Note: ").bold = True
    p.add_run(text)
    p.add_run(" — see appendix for details.").italic = True

# ---------------------------------------------------------------------------
# Chapter 1: Introduction
# ---------------------------------------------------------------------------
doc.add_heading("Chapter 1: Introduction to Software Engineering", 1)
doc.add_heading("1.1 Overview", 2)
lorem(doc, 3)
add_mixed_paragraph(doc, "Software engineering encompasses a wide range of disciplines "
    "from requirements analysis to deployment and maintenance.")

doc.add_heading("1.2 Historical Context", 2)
lorem(doc, 2)
add_bullet_list(doc, [
    "1960s: Birth of structured programming",
    "1970s: Software crisis and the rise of methodologies",
    "1980s: Object-oriented programming emerges",
    "1990s: Agile manifesto and iterative development",
    "2000s: DevOps, cloud computing, microservices",
    "2010s: AI/ML integration in software workflows",
    "2020s: LLM-assisted development",
])

doc.add_heading("1.3 Core Principles", 2)
lorem(doc, 2)
add_numbered_list(doc, [
    "Separation of concerns",
    "DRY (Don't Repeat Yourself)",
    "SOLID principles",
    "Fail fast, fail loudly",
    "Composability over inheritance",
])

add_table(doc,
    ["Principle", "Description", "Example"],
    [
        ["SRP", "Single Responsibility Principle", "One class per concern"],
        ["OCP", "Open/Closed Principle", "Extend, don't modify"],
        ["LSP", "Liskov Substitution Principle", "Subtypes are substitutable"],
        ["ISP", "Interface Segregation", "Many small interfaces"],
        ["DIP", "Dependency Inversion", "Depend on abstractions"],
    ])
doc.add_page_break()

# ---------------------------------------------------------------------------
# Chapter 2: Requirements Engineering
# ---------------------------------------------------------------------------
doc.add_heading("Chapter 2: Requirements Engineering", 1)
doc.add_heading("2.1 Elicitation Techniques", 2)
lorem(doc, 3)

doc.add_heading("2.1.1 Interviews", 3)
lorem(doc, 2)
add_bullet_list(doc, [
    "Structured interviews: fixed questions, quantitative data",
    "Semi-structured: guided conversation with flexibility",
    "Unstructured: open exploration of stakeholder needs",
])

doc.add_heading("2.1.2 Workshops", 3)
lorem(doc, 2)

doc.add_heading("2.2 Specification Formats", 2)
lorem(doc, 2)

add_table(doc,
    ["Format", "Formality", "Use Case", "Tooling"],
    [
        ["User Stories", "Low", "Agile sprints", "Jira, Linear"],
        ["Use Cases", "Medium", "UML modeling", "Enterprise Architect"],
        ["SRS Document", "High", "Regulated industries", "Confluence, Word"],
        ["BDD Scenarios", "Medium", "Test-driven", "Cucumber, Behave"],
    ])

doc.add_heading("2.3 Acceptance Criteria", 2)
lorem(doc, 3)
add_mixed_paragraph(doc, "Acceptance criteria must be measurable, verifiable, and unambiguous.")
doc.add_page_break()

# ---------------------------------------------------------------------------
# Chapter 3: System Design
# ---------------------------------------------------------------------------
doc.add_heading("Chapter 3: System Design", 1)
doc.add_heading("3.1 Architectural Patterns", 2)
lorem(doc, 2)

doc.add_heading("3.1.1 Monolithic Architecture", 3)
lorem(doc, 2)
p = doc.add_paragraph()
p.add_run("Advantages: ").bold = True
p.add_run("Simple deployment, easy debugging, low operational overhead.")
p = doc.add_paragraph()
p.add_run("Disadvantages: ").bold = True
p.add_run("Tight coupling, difficult to scale horizontally, long build times.")

doc.add_heading("3.1.2 Microservices", 3)
lorem(doc, 2)
add_bullet_list(doc, [
    "Independent deployability per service",
    "Polyglot persistence (each service owns its data store)",
    "Failure isolation through circuit breakers",
    "Service mesh for traffic management (Istio, Linkerd)",
])

doc.add_heading("3.1.3 Event-Driven Architecture", 3)
lorem(doc, 2)
add_table(doc,
    ["Pattern", "Broker", "Use Case"],
    [
        ["Pub/Sub", "Kafka, Pub/Sub", "Stream processing"],
        ["Event Sourcing", "EventStore", "Audit trail, CQRS"],
        ["Saga", "Conductor", "Distributed transactions"],
        ["Outbox Pattern", "Debezium", "Reliable messaging"],
    ])

doc.add_heading("3.2 Data Modeling", 2)
lorem(doc, 3)
doc.add_page_break()

# ---------------------------------------------------------------------------
# Chapter 4: Development Practices
# ---------------------------------------------------------------------------
doc.add_heading("Chapter 4: Development Practices", 1)
doc.add_heading("4.1 Version Control Strategies", 2)
lorem(doc, 2)

doc.add_heading("4.1.1 Branching Models", 3)
add_numbered_list(doc, [
    "Trunk-based development: single long-lived main branch",
    "Git Flow: feature/release/hotfix branch model",
    "GitHub Flow: simplified flow with feature branches and PRs",
    "GitLab Flow: environment-based branching",
])
lorem(doc, 2)

doc.add_heading("4.2 Code Review", 2)
lorem(doc, 2)
add_table(doc,
    ["Practice", "Goal", "Anti-Pattern"],
    [
        ["Pair Review", "Knowledge sharing", "Rubber-stamping"],
        ["Automated Checks", "Consistency", "Gate-keeping"],
        ["Author Self-Review", "Catch obvious bugs", "Skipping"],
        ["Async Reviews", "Parallel work", "Long-pending PRs"],
    ])

doc.add_heading("4.3 Testing Pyramid", 2)
lorem(doc, 2)
p = doc.add_paragraph()
p.add_run("Unit tests ").bold = True
p.add_run("form the base: fast, isolated, cheap. ")
p.add_run("Integration tests ").bold = True
p.add_run("verify component interaction. ")
p.add_run("End-to-end tests ").bold = True
p.add_run("validate user journeys but are slow and expensive.")

add_numbered_list(doc, [
    "Unit: 70% of test suite — sub-millisecond execution",
    "Integration: 20% — test real dependencies (DB, queues)",
    "E2E: 10% — critical user paths only",
])
doc.add_page_break()

# ---------------------------------------------------------------------------
# Chapter 5: DevOps & CI/CD
# ---------------------------------------------------------------------------
doc.add_heading("Chapter 5: DevOps and Continuous Delivery", 1)
doc.add_heading("5.1 CI Pipeline Design", 2)
lorem(doc, 2)
add_bullet_list(doc, [
    "Lint and format check (ruff, ESLint, etc.)",
    "Static type checking (mypy, TypeScript)",
    "Unit test execution with coverage gate",
    "Build artifact (Docker image, wheel, binary)",
    "Integration test against ephemeral environment",
    "Security scan (Trivy, Snyk, Dependabot)",
    "Publish to staging registry",
])

doc.add_heading("5.2 Deployment Strategies", 2)
lorem(doc, 2)
add_table(doc,
    ["Strategy", "Rollout", "Rollback", "Downtime"],
    [
        ["Big Bang", "Immediate", "Manual restore", "Yes"],
        ["Blue/Green", "Switch traffic", "Switch back", "No"],
        ["Canary", "Gradual %", "Reduce %", "No"],
        ["Shadow", "Duplicate traffic", "Remove shadow", "No"],
        ["Rolling", "Pod-by-pod", "Version revert", "Minimal"],
    ])

doc.add_heading("5.3 Observability", 2)
lorem(doc, 2)

doc.add_heading("5.3.1 The Three Pillars", 3)
add_numbered_list(doc, [
    "Logs: structured, searchable event records",
    "Metrics: time-series aggregates (Prometheus, Datadog)",
    "Traces: distributed request span correlation (OpenTelemetry)",
])
lorem(doc, 2)
doc.add_page_break()

# ---------------------------------------------------------------------------
# Chapter 6: Security Engineering
# ---------------------------------------------------------------------------
doc.add_heading("Chapter 6: Security Engineering", 1)
doc.add_heading("6.1 OWASP Top 10", 2)
lorem(doc, 2)
add_table(doc,
    ["#", "Vulnerability", "Mitigation"],
    [
        ["A01", "Broken Access Control", "RBAC, least privilege"],
        ["A02", "Cryptographic Failures", "TLS 1.3, modern ciphers"],
        ["A03", "Injection", "Parameterized queries, input validation"],
        ["A04", "Insecure Design", "Threat modeling, secure by design"],
        ["A05", "Security Misconfiguration", "Hardened defaults, IaC"],
        ["A06", "Vulnerable Components", "Dependency scanning, SCA"],
        ["A07", "Auth Failures", "MFA, secure session management"],
        ["A08", "Software Integrity Failures", "Supply chain verification"],
        ["A09", "Logging Failures", "Centralized SIEM, audit trails"],
        ["A10", "SSRF", "Allowlist outbound connections"],
    ])

doc.add_heading("6.2 Secure Development Lifecycle", 2)
lorem(doc, 3)
add_bullet_list(doc, [
    "Threat modeling at design phase (STRIDE, PASTA)",
    "SAST: static analysis before merge (CodeQL, Semgrep)",
    "DAST: dynamic scanning in staging (OWASP ZAP, Burp Suite)",
    "Penetration testing annually or after major releases",
    "Security champions program — embed sec in dev teams",
])

doc.add_heading("6.3 Secrets Management", 2)
lorem(doc, 2)
p = doc.add_paragraph()
p.add_run("Never ").bold = True
p.add_run("store secrets in source code. Use a secrets manager: ")
p.add_run("HashiCorp Vault, AWS Secrets Manager, GCP Secret Manager").italic = True
p.add_run(". Rotate secrets regularly and audit access logs.")
doc.add_page_break()

# ---------------------------------------------------------------------------
# Chapter 7: Performance Engineering
# ---------------------------------------------------------------------------
doc.add_heading("Chapter 7: Performance Engineering", 1)
doc.add_heading("7.1 Performance Testing Types", 2)
lorem(doc, 2)
add_table(doc,
    ["Type", "Goal", "Tool"],
    [
        ["Load Test", "Verify at expected load", "k6, JMeter, Locust"],
        ["Stress Test", "Find breaking point", "k6, Gatling"],
        ["Soak Test", "Detect memory leaks", "Grafana k6"],
        ["Spike Test", "Handle sudden traffic", "k6, Artillery"],
        ["Chaos Test", "Failure resilience", "Chaos Monkey, Gremlin"],
    ])

doc.add_heading("7.2 Profiling and Optimization", 2)
lorem(doc, 3)
add_numbered_list(doc, [
    "Profile first, optimize second — never guess",
    "Database query optimization: explain plans, index design",
    "Caching strategy: CDN, application cache, DB query cache",
    "Async processing: offload heavy work to queues",
    "Horizontal scaling: stateless services behind load balancer",
    "Connection pooling: reuse DB connections",
])
doc.add_page_break()

# ---------------------------------------------------------------------------
# Appendix A: Glossary
# ---------------------------------------------------------------------------
doc.add_heading("Appendix A: Glossary", 1)
lorem(doc)
add_table(doc,
    ["Term", "Definition"],
    [
        ["API", "Application Programming Interface"],
        ["CI/CD", "Continuous Integration / Continuous Delivery"],
        ["CQRS", "Command Query Responsibility Segregation"],
        ["DDD", "Domain-Driven Design"],
        ["IaC", "Infrastructure as Code"],
        ["MTTR", "Mean Time To Recovery"],
        ["SLA", "Service Level Agreement"],
        ["SLI", "Service Level Indicator"],
        ["SLO", "Service Level Objective"],
        ["TTL", "Time To Live"],
    ])

# ---------------------------------------------------------------------------
# Appendix B: Bibliography
# ---------------------------------------------------------------------------
doc.add_heading("Appendix B: Bibliography", 1)
add_numbered_list(doc, [
    "Martin, Robert C. Clean Code. Prentice Hall, 2008.",
    "Fowler, Martin. Refactoring. Addison-Wesley, 2018.",
    "Newman, Sam. Building Microservices. O'Reilly, 2021.",
    "Kim, Gene et al. The DevOps Handbook. IT Revolution Press, 2016.",
    "Evans, Eric. Domain-Driven Design. Addison-Wesley, 2003.",
    "OWASP Foundation. OWASP Top 10 2021. https://owasp.org/Top10/",
    "Google SRE Team. Site Reliability Engineering. O'Reilly, 2016.",
])
lorem(doc, 2)

# Save
doc.save("test_report.docx")
print("Generated test_report.docx")
